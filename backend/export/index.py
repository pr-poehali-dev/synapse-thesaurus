import json
import os
from typing import Dict, Any, List
from io import BytesIO
import base64
from datetime import datetime

def create_pdf(text: str, replacements: List[Dict]) -> bytes:
    """
    Создает PDF с текстом и историей замен
    Использует простой подход без библиотек
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#0EA5E9'),
            spaceAfter=30
        )
        
        story.append(Paragraph("Synapse - Exported Document", title_style))
        story.append(Spacer(1, 0.2 * inch))
        
        story.append(Paragraph("Main Text", styles['Heading2']))
        story.append(Spacer(1, 0.1 * inch))
        
        text_paragraphs = text.split('\n')
        for para in text_paragraphs:
            if para.strip():
                story.append(Paragraph(para, styles['BodyText']))
                story.append(Spacer(1, 0.1 * inch))
        
        if replacements:
            story.append(Spacer(1, 0.3 * inch))
            story.append(Paragraph("Replacement History", styles['Heading2']))
            story.append(Spacer(1, 0.1 * inch))
            
            table_data = [['Original', 'Replacement', 'Timestamp']]
            for rep in replacements:
                table_data.append([
                    rep.get('original', ''),
                    rep.get('replacement', ''),
                    rep.get('timestamp', '')
                ])
            
            t = Table(table_data, colWidths=[2*inch, 2*inch, 2*inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0EA5E9')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(t)
        
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes
        
    except ImportError:
        return create_simple_pdf_fallback(text, replacements)

def create_simple_pdf_fallback(text: str, replacements: List[Dict]) -> bytes:
    """
    Простой текстовый документ если reportlab недоступен
    """
    content = f"""SYNAPSE - EXPORTED DOCUMENT
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

MAIN TEXT:
{text}

"""
    if replacements:
        content += "\nREPLACEMENT HISTORY:\n"
        for rep in replacements:
            content += f"  {rep.get('original', '')} → {rep.get('replacement', '')} ({rep.get('timestamp', '')})\n"
    
    return content.encode('utf-8')

def create_docx(text: str, replacements: List[Dict]) -> bytes:
    """
    Создает DOCX файл с текстом и историей замен
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        title = doc.add_heading('Synapse - Exported Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(14, 165, 233)
        
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        doc.add_heading('Main Text', level=1)
        
        text_paragraphs = text.split('\n')
        for para in text_paragraphs:
            if para.strip():
                p = doc.add_paragraph(para)
                p.style = 'Normal'
        
        if replacements:
            doc.add_page_break()
            doc.add_heading('Replacement History', level=1)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Original'
            hdr_cells[1].text = 'Replacement'
            hdr_cells[2].text = 'Timestamp'
            
            for rep in replacements:
                row_cells = table.add_row().cells
                row_cells[0].text = rep.get('original', '')
                row_cells[1].text = rep.get('replacement', '')
                row_cells[2].text = rep.get('timestamp', '')
        
        buffer = BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes
        
    except ImportError:
        return create_simple_pdf_fallback(text, replacements)

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    """
    Экспорт текста в PDF или DOCX формат
    Args: event - dict с httpMethod, body (text, replacements, format)
    Returns: HTTP response с base64 файлом
    """
    method: str = event.get('httpMethod', 'POST')
    
    if method == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Max-Age': '86400'
            },
            'body': '',
            'isBase64Encoded': False
        }
    
    if method == 'POST':
        try:
            body_data = json.loads(event.get('body', '{}'))
            text = body_data.get('text', '')
            replacements = body_data.get('replacements', [])
            export_format = body_data.get('format', 'pdf').lower()
            
            if not text:
                return {
                    'statusCode': 400,
                    'headers': {
                        'Content-Type': 'application/json',
                        'Access-Control-Allow-Origin': '*'
                    },
                    'body': json.dumps({'error': 'Text is required'}),
                    'isBase64Encoded': False
                }
            
            if export_format == 'pdf':
                file_bytes = create_pdf(text, replacements)
                content_type = 'application/pdf'
                filename = f'synapse-export-{datetime.now().strftime("%Y%m%d-%H%M%S")}.pdf'
            elif export_format == 'docx':
                file_bytes = create_docx(text, replacements)
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                filename = f'synapse-export-{datetime.now().strftime("%Y%m%d-%H%M%S")}.docx'
            else:
                return {
                    'statusCode': 400,
                    'headers': {
                        'Content-Type': 'application/json',
                        'Access-Control-Allow-Origin': '*'
                    },
                    'body': json.dumps({'error': 'Invalid format. Use pdf or docx'}),
                    'isBase64Encoded': False
                }
            
            file_base64 = base64.b64encode(file_bytes).decode('utf-8')
            
            return {
                'statusCode': 200,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*'
                },
                'body': json.dumps({
                    'filename': filename,
                    'content': file_base64,
                    'contentType': content_type,
                    'size': len(file_bytes)
                }),
                'isBase64Encoded': False
            }
            
        except Exception as e:
            return {
                'statusCode': 500,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*'
                },
                'body': json.dumps({'error': str(e)}),
                'isBase64Encoded': False
            }
    
    return {
        'statusCode': 405,
        'headers': {
            'Content-Type': 'application/json',
            'Access-Control-Allow-Origin': '*'
        },
        'body': json.dumps({'error': 'Method not allowed'}),
        'isBase64Encoded': False
    }
